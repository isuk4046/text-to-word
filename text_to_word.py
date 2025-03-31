import os
import sys
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_text_to_docx(input_file, output_file=None):
    """
    Convert a text file to a Word document (.docx)
    
    Args:
        input_file (str): Path to the input text file
        output_file (str, optional): Path to the output Word file. 
                                     If None, uses the same name as input with .docx extension
    
    Returns:
        str: Path to the created Word document or None if conversion failed
    """
    # Validate input file
    if not os.path.isfile(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        return None
    
    # Create output filename if not specified
    if output_file is None:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}.docx"
    
    try:
        # Read text file content
        with open(input_file, 'r', encoding='utf-8') as file:
            text_content = file.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Process text line by line for better formatting
        for paragraph_text in text_content.split('\n'):
            if paragraph_text.strip() == '':
                # Add empty paragraph for blank lines
                doc.add_paragraph()
            else:
                # Add paragraph with text
                doc.add_paragraph(paragraph_text)
        
        # Save the document
        doc.save(output_file)
        print(f"Successfully converted '{input_file}' to '{output_file}'")
        return output_file
    
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        return None

def batch_convert(input_files, output_dir=None, output_files=None):
    """
    Convert multiple text files to Word documents
    
    Args:
        input_files (list): List of input file paths or glob patterns
        output_dir (str, optional): Directory for output files
        output_files (list, optional): List of output file paths (must match input_files length)
    
    Returns:
        list: Paths to successfully created Word documents
    """
    # Expand any glob patterns in input_files
    expanded_input_files = []
    for file_pattern in input_files:
        matches = glob.glob(file_pattern)
        if matches:
            expanded_input_files.extend(matches)
        else:
            print(f"Warning: No files match pattern '{file_pattern}'")
    
    if not expanded_input_files:
        print("No valid input files found.")
        return []
    
    # Create output directory if specified and doesn't exist
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Process each file
    successful_conversions = []
    for i, input_file in enumerate(expanded_input_files):
        # Determine output file path
        if output_files and i < len(output_files):
            # Use specified output filename
            output_file = output_files[i]
        else:
            # Generate output filename based on input
            base_name = os.path.basename(input_file)
            name_without_ext = os.path.splitext(base_name)[0]
            if output_dir:
                output_file = os.path.join(output_dir, f"{name_without_ext}.docx")
            else:
                output_file = f"{os.path.splitext(input_file)[0]}.docx"
        
        # Convert the file
        result = convert_text_to_docx(input_file, output_file)
        if result:
            successful_conversions.append(result)
    
    print(f"\nConversion complete: {len(successful_conversions)} of {len(expanded_input_files)} files converted successfully.")
    return successful_conversions

def main():
    """Command line interface for the converter"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert text files to Word documents (.docx)')
    parser.add_argument('input', nargs='+', help='Input text file(s) or glob patterns')
    parser.add_argument('-o', '--output', nargs='*', help='Output file name(s)')
    parser.add_argument('-d', '--output-dir', help='Output directory for all files')
    
    args = parser.parse_args()
    
    # Validate output files if specified
    if args.output and len(args.output) > 0 and len(args.output) != len(args.input):
        print("Error: If specifying output files, the number must match the number of input files.")
        return
    
    batch_convert(args.input, args.output_dir, args.output)

if __name__ == "__main__":
    main()